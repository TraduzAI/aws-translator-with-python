# services/document_service.py

"""
Document Service Module
=======================

Este módulo fornece serviços para importar e exportar textos a partir e para
diferentes formatos de documentos. Suporta os seguintes formatos:

- **Importação**:
    - PDF (`.pdf`)
    - DOCX (`.docx`)
    - EPUB (`.epub`)
    - TXT (`.txt`)
- **Exportação**:
    - PDF (`.pdf`)
    - DOCX (`.docx`)
    - TXT (`.txt`)

Classes:
    DocumentService: Classe responsável pela importação e exportação de documentos.

Dependências:
    - PyPDF2: biblioteca para manipulação de arquivos PDF.
    - python-docx: biblioteca para manipulação de arquivos DOCX.
    - EbookLib: biblioteca para manipulação de arquivos EPUB.
    - reportlab: biblioteca para geração de PDFs.
    - typing: biblioteca padrão para anotações de tipos.

Exemplo de Uso:
    >>> from services.document_service import DocumentService
    >>> doc_service = DocumentService()

    # Importar um documento
    >>> texto = doc_service.import_document('exemplo.pdf')

    # Exportar um documento
    >>> doc_service.export_document(texto, 'saida.docx', 'docx')
"""

from typing import Optional
import os

import PyPDF2  # Para PDFs
from docx import Document  # Para DOCX
from ebooklib import epub  # Para EPUB

from reportlab.lib.pagesizes import letter  # Para exportar PDFs
from reportlab.pdfgen import canvas


class DocumentService:
    """
    Serviço para importar e exportar textos a partir e para diferentes formatos de documentos.

    Esta classe fornece métodos para importar textos de arquivos PDF, DOCX, EPUB e TXT, bem
    como exportar textos para arquivos PDF, DOCX e TXT. Utiliza diversas bibliotecas para
    manipulação de diferentes formatos de arquivos.

    Métodos:
        import_document(file_path: str) ⇾ Optional[str]:
            Importa texto de um arquivo de documento.

        export_document(text: str, file_path: str, format: str, metrics_original: dict = None,
                        metrics_simplified: dict = None, bleu_score: float = None) ⇾ None:
            Exporta texto para um arquivo de documento, incluindo o BLEU Score.
    """

    def import_document(self, file_path: str) -> Optional[str]:
        """
        Importa texto de um arquivo de documento.

        Este metodo determina o tipo de arquivo com base na extensão e utiliza o metodo
        apropriado para extrair o texto.

        Parâmetros:
            file_path (str): Caminho para o arquivo de documento a ser importado.

        Retorna:
            Optional[str]: O texto extraído do documento ou `None` se não for possível extrair.

        Exceções:
            - ValueError: se o formato do arquivo não for suportado.
            - Exception: Se ocorrer um erro durante a importação do documento.

        Exemplos de Uso:
            >>> doc_service = DocumentService()
            >>> texto = doc_service.import_document('documento.pdf')
            >>> print(texto)
            "Este é o conteúdo extraído do PDF."
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        if ext == '.pdf':
            return self._import_pdf(file_path)
        elif ext == '.docx':
            return self._import_docx(file_path)
        elif ext == '.epub':
            return self._import_epub(file_path)
        elif ext == '.txt':
            return self._import_txt(file_path)
        else:
            raise ValueError(f"Formato de arquivo não suportado: {ext}")

    def export_document(self, text: str, file_path: str, format: str, metrics_original: dict = None,
                        metrics_simplified: dict = None, bleu_score: float = None) -> None:
        """
        Exporta texto e métricas para um arquivo de documento, incluindo o BLEU Score.

        Este metodo determina o formato de exportação com base no parâmetro `format`
        e utiliza o metodo apropriado para salvar o texto, métricas e BLEU Score no formato desejado.

        Parâmetros:
            text (str): O texto a ser exportado.
            file_path (str): Caminho onde o arquivo exportado será salvo.
            format (str): Formato de exportação desejado (`'pdf'`, `'docx'`, `'txt'`).
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.
            bleu_score (float): O BLEU Score do texto simplificado e traduzido.

        Retorna:
            None

        Exceções:
            - ValueError: se o formato de exportação não for suportado.
            - Exception: Se ocorrer um erro durante a exportação do documento.
        """
        format = format.lower()
        if format == 'pdf':
            self._export_pdf(text, file_path, metrics_original, metrics_simplified, bleu_score)
        elif format == 'docx':
            self._export_docx(text, file_path, metrics_original, metrics_simplified, bleu_score)
        elif format == 'txt':
            self._export_txt(text, file_path, metrics_original, metrics_simplified, bleu_score)
        else:
            raise ValueError(f"Formato de exportação não suportado: {format}")

    @staticmethod
    def _import_pdf(file_path: str) -> str:
        """
        Importa texto de um arquivo PDF.

        Utiliza a biblioteca PyPDF2 para extrair o texto de cada página do PDF.

        Parâmetros:
            file_path (str): Caminho para o arquivo PDF a ser importado.

        Retorna:
            str: O texto extraído do PDF.

        Exceções:
            - FileNotFoundError: se o arquivo PDF não for encontrado.
            - Exception: Se ocorrer um erro durante a leitura do PDF.
        """
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''
                for page in reader.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + '\n'
                return text.strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"Arquivo PDF não encontrado: {file_path}")
        except Exception as e:
            raise Exception(f"Erro ao importar PDF: {str(e)}")

    @staticmethod
    def _import_docx(file_path: str) -> str:
        """
        Importa texto de um arquivo DOCX.

        Utiliza a biblioteca python-docx para extrair o texto de cada parágrafo do DOCX.

        Parâmetros:
            file_path (str): Caminho para o arquivo DOCX a ser importado.

        Retorna:
            str: O texto extraído do DOCX.

        Exceções:
            - FileNotFoundError: se o arquivo DOCX não for encontrado.
            - Exception: Se ocorrer um erro durante a leitura do DOCX.
        """
        try:
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text.strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"Arquivo DOCX não encontrado: {file_path}")
        except Exception as e:
            raise Exception(f"Erro ao importar DOCX: {str(e)}")

    @staticmethod
    def _import_epub(file_path: str) -> str:
        """
        Importa texto de um arquivo EPUB.

        Utiliza a biblioteca EbookLib para extrair o conteúdo textual dos documentos do EPUB.

        Parâmetros:
            file_path (str): Caminho para o arquivo EPUB a ser importado.

        Retorna:
            str: O texto extraído do EPUB.

        Exceções:
            - FileNotFoundError: se o arquivo EPUB não for encontrado.
            - Exception: Se ocorrer um erro durante a leitura do EPUB.
        """
        try:
            book = epub.read_epub(file_path)
            text = ''
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    content = item.get_content()
                    text += content.decode('utf-8') + '\n'
            return text.strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"Arquivo EPUB não encontrado: {file_path}")
        except Exception as e:
            raise Exception(f"Erro ao importar EPUB: {str(e)}")

    @staticmethod
    def _import_txt(file_path: str) -> str:
        """
        Importa texto de um arquivo TXT.

        Abre o arquivo de texto e lê todos o seu conteúdo.

        Parâmetros:
            file_path (str): Caminho para o arquivo TXT a ser importado.

        Retorna:
            str: O texto extraído do TXT.

        Exceções:
            - FileNotFoundError: se o arquivo TXT não for encontrado.
            - Exception: Se ocorrer um erro durante a leitura do TXT.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"Arquivo TXT não encontrado: {file_path}")
        except Exception as e:
            raise Exception(f"Erro ao importar TXT: {str(e)}")

    @staticmethod
    def _export_pdf(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None,
                    bleu_score: float = None) -> None:
        """
        Exporta texto e métricas para um arquivo PDF, incluindo o BLEU Score.

        Utiliza a biblioteca ReportLab para gerar um PDF a partir do texto fornecido,
        cuidando da formatação e quebra de linhas conforme necessário, e inclui as métricas e o BLEU Score.

        Parâmetros:
            text (str): O texto a ser exportado para o PDF.
            file_path (str): Caminho onde o arquivo PDF será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.
            bleu_score (float): O BLEU Score do texto simplificado e traduzido.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a criação do PDF.
        """
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            # Configurações do texto
            text_object = c.beginText(50, height - 50)
            text_object.setFont("Helvetica-Bold", 14)
            text_object.textLine("Texto Simplificado e Traduzido:")
            text_object.setFont("Helvetica", 12)
            text_object.textLine("")

            # Adicionar o texto
            for line in text.split('\n'):
                words = line.split(' ')
                line_buffer = ""

                for word in words:
                    if c.stringWidth(line_buffer + word, "Helvetica", 12) < (width - 100):
                        line_buffer += word + " "
                    else:
                        text_object.textLine(line_buffer.strip())
                        line_buffer = word + " "

                        if text_object.getY() <= 50:
                            c.drawText(text_object)
                            c.showPage()
                            text_object = c.beginText(50, height - 50)
                            text_object.setFont("Helvetica", 12)

                if line_buffer:
                    text_object.textLine(line_buffer.strip())

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

            text_object.textLine("")

            if metrics_original and metrics_simplified:
                metric_names = {
                    'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                    'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                    'smog_index': 'Índice SMOG',
                    'coleman_liau_index': 'Índice de Coleman-Liau',
                    'automated_readability_index': 'Índice ARI',
                    'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                }

                # Métricas do texto original
                text_object.setFont("Helvetica-Bold", 14)
                text_object.textLine("Métricas do Texto Original:")
                text_object.setFont("Helvetica", 12)
                text_object.textLine("")

                for key, value in metrics_original.items():
                    metric_name = metric_names.get(key, key)
                    text_object.textLine(f"{metric_name}: {value:.2f}")

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

                text_object.textLine("")

                # Métricas do texto simplificado
                text_object.setFont("Helvetica-Bold", 14)
                text_object.textLine("Métricas do Texto Simplificado:")
                text_object.setFont("Helvetica", 12)
                text_object.textLine("")

                for key, value in metrics_simplified.items():
                    metric_name = metric_names.get(key, key)
                    text_object.textLine(f"{metric_name}: {value:.2f}")

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

            # Adicionar o BLEU Score
            if bleu_score is not None:
                text_object.textLine("")
                text_object.setFont("Helvetica-Bold", 14)
                text_object.textLine("BLEU Score:")
                text_object.setFont("Helvetica", 12)
                text_object.textLine(f"{bleu_score:.2f}")

                if text_object.getY() <= 50:
                    c.drawText(text_object)
                    c.showPage()
                    text_object = c.beginText(50, height - 50)
                    text_object.setFont("Helvetica", 12)

            c.drawText(text_object)
            c.save()
        except Exception as e:
            raise Exception(f"Erro ao exportar PDF: {str(e)}")

    @staticmethod
    def _export_docx(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None,
                     bleu_score: float = None) -> None:
        """
        Exporta texto e métricas para um arquivo DOCX, incluindo o BLEU Score.

        Utiliza a biblioteca python-docx para criar um documento DOCX com o texto e as métricas fornecidos.

        Parâmetros:
            text (str): O texto a ser exportado para o DOCX.
            file_path (str): Caminho onde o arquivo DOCX será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.
            bleu_score (float): O BLEU Score do texto simplificado e traduzido.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a criação do DOCX.
        """
        try:
            doc = Document()

            # Adicionar título
            doc.add_heading('Texto Simplificado e Traduzido:', level=1)
            doc.add_paragraph(text)

            if metrics_original and metrics_simplified:
                metric_names = {
                    'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                    'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                    'smog_index': 'Índice SMOG',
                    'coleman_liau_index': 'Índice de Coleman-Liau',
                    'automated_readability_index': 'Índice ARI',
                    'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                }

                # Métricas do texto original
                doc.add_heading('Métricas do Texto Original:', level=2)
                for key, value in metrics_original.items():
                    metric_name = metric_names.get(key, key)
                    doc.add_paragraph(f"{metric_name}: {value:.2f}")

                # Métricas do texto simplificado
                doc.add_heading('Métricas do Texto Simplificado:', level=2)
                for key, value in metrics_simplified.items():
                    metric_name = metric_names.get(key, key)
                    doc.add_paragraph(f"{metric_name}: {value:.2f}")

            # Adicionar o BLEU Score
            if bleu_score is not None:
                doc.add_heading('BLEU Score:', level=2)
                doc.add_paragraph(f"{bleu_score:.2f}")

            doc.save(file_path)
        except Exception as e:
            raise Exception(f"Erro ao exportar DOCX: {str(e)}")

    @staticmethod
    def _export_txt(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None,
                    bleu_score: float = None) -> None:
        """
        Exporta texto e métricas para um arquivo TXT, incluindo o BLEU Score.

        Abre (ou cria) o arquivo de texto e escreve todos os conteúdos fornecidos, incluindo as métricas e o BLEU Score.

        Parâmetros:
            text (str): O texto a ser exportado para o TXT.
            file_path (str): Caminho onde o arquivo TXT será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.
            bleu_score (float): O BLEU Score do texto simplificado e traduzido.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a escrita no TXT.
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("Texto Simplificado e Traduzido:\n")
                f.write(text)
                f.write("\n\n")

                if metrics_original and metrics_simplified:
                    metric_names = {
                        'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                        'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                        'smog_index': 'Índice SMOG',
                        'coleman_liau_index': 'Índice de Coleman-Liau',
                        'automated_readability_index': 'Índice ARI',
                        'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                    }

                    f.write("Métricas do Texto Original:\n")
                    for key, value in metrics_original.items():
                        metric_name = metric_names.get(key, key)
                        f.write(f"{metric_name}: {value:.2f}\n")

                    f.write("\nMétricas do Texto Simplificado:\n")
                    for key, value in metrics_simplified.items():
                        metric_name = metric_names.get(key, key)
                        f.write(f"{metric_name}: {value:.2f}\n")

                # Adicionar o BLEU Score
                if bleu_score is not None:
                    f.write("\nBLEU Score:\n")
                    f.write(f"{bleu_score:.2f}\n")
        except Exception as e:
            raise Exception(f"Erro ao exportar TXT: {str(e)}")
